"""
知识库构建脚本 - 纯Python实现
数据加载 -> 清洗 -> 分块 -> TF-IDF向量化 -> 存储

支持 PDF/DOCX 文件，使用 TF-IDF 嵌入 + 余弦相似度检索
无任何原生二进制依赖（不依赖 chromadb/onnxruntime）
"""

import os
import re
import json
import hashlib
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from collections import Counter
import math

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)

# ─── 配置常量 ──────────────────────────────────────
DATA_DIR = Path(__file__).parent / "河北大学规章制度"
VECTOR_DB_DIR = Path(__file__).parent / "vector_db"
COLLECTION_NAME = "campus_rules"
CHUNK_SIZE = 600       # 增大分块，确保每个片段包含完整条款
CHUNK_OVERLAP = 100    # 增大重叠，保证跨块信息不丢失


# ══════════════════════════════════════════════════
# 1. 文档加载
# ══════════════════════════════════════════════════

def load_pdf(filepath: str) -> str:
    """从PDF文件提取文本"""
    try:
        import pdfplumber
        texts: list[str] = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    texts.append(text.strip())
        return "\n".join(texts)
    except Exception as e:
        logger.error(f"PDF 加载失败 {filepath}: {e}")
        return ""


def load_docx(filepath: str) -> str:
    """从 DOCX 文件提取文本"""
    try:
        from docx import Document
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.error(f"DOCX 加载失败 {filepath}: {e}")
        return ""


def load_document(filepath: str) -> str:
    """根据扩展名自动选择加载器"""
    ext = Path(filepath).suffix.lower()
    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".doc": load_docx,
    }
    loader = loaders.get(ext)
    if loader is None:
        logger.warning(f"不支持的格式: {ext} ({filepath})")
        return ""
    return loader(filepath)


# ══════════════════════════════════════════════════
# 2. 文本清洗
# ══════════════════════════════════════════════════

def clean_text(text: str) -> str:
    """
    文本清洗：
    - 合并多余空白、换行
    - 去除特殊控制字符
    - 保留中文标点
    """
    # 替换各种空白字符为空格
    text = re.sub(r"[\t\f\r\v]+", " ", text)
    # 合并连续空行（保留段落分隔）
    text = re.sub(r"\n{3,}", "\n\n", text)
    # 去除首尾空白
    text = text.strip()
    # 去除不可见控制字符（0x00-0x08, 0x0B, 0x0C, 0x0E-0x1F），但保留中文
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    # 去除页眉页脚等常见噪音模式
    text = re.sub(r"第\s*\d+\s*页", "", text)
    return text


# ══════════════════════════════════════════════════
# 3. 文本分块
# ══════════════════════════════════════════════════

def split_text_into_paragraphs(text: str) -> List[str]:
    """智能分段：尝试多种分隔符，确保长文档能被正确分割"""
    
    # 策略1：双换行（标准段落）
    parts = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    # 如果分段后大部分段落长度合理（< 500字），就用这种分割
    if parts:
        long_count = sum(1 for p in parts if len(p) > 500)
        if long_count <= len(parts) * 0.5:  # 不超过一半是超长段落
            return parts
    
    # 策略2：单换行
    parts = [p.strip() for p in text.split('\n') if len(p.strip()) > 10]
    if len(parts) >= 3:
        return parts
    
    # 策略3：按中文句号/分号/英文句号分割
    parts = re.split(r'(?<=[。；.!;])\s*', text)
    return [p.strip() for p in parts if len(p.strip()) >= 10]


def split_documents(documents: List[Any], chunk_size: int = CHUNK_SIZE,
                   overlap: int = CHUNK_OVERLAP) -> List[Any]:
    """
    将长文档按段落智能分割为固定大小的文本块
    尽量在句号/段落边界处切分
    """
    from langchain_core.documents import Document as LCDocument

    chunks = []
    for doc in documents:
        text = doc.page_content
        if not text or len(text) < 10:
            continue

        metadata = dict(doc.metadata)

        if len(text) <= chunk_size:
            chunks.append(LCDocument(page_content=text, metadata=metadata))
            continue

        # 使用增强的段分割函数
        paragraphs = split_text_into_paragraphs(text)
        current_chunk = ""
        current_len = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            para_len = len(para)

            if current_len == 0:
                current_chunk = para
                current_len = para_len
            elif current_len + para_len + 1 <= chunk_size:
                current_chunk += "\n" + para
                current_len += para_len + 1
            else:
                chunks.append(LCDocument(page_content=current_chunk.strip(), metadata=metadata))
                if overlap > 0 and len(current_chunk) > overlap:
                    overlap_text = current_chunk[-overlap:]
                    last_nl = overlap_text.find("\n")
                    if last_nl >= 0:
                        current_chunk = overlap_text[last_nl + 1:] + "\n" + para
                    else:
                        current_chunk = current_chunk[-overlap:] + "\n" + para
                else:
                    current_chunk = para
                current_len = len(current_chunk)

        if current_chunk.strip():
            chunks.append(LCDocument(page_content=current_chunk.strip(), metadata=metadata))

    return chunks


# ══════════════════════════════════════════════════
# 4. TF-IDF 嵌入 (纯 Python + NumPy)
# ══════════════════════════════════════════════════

class TfidfVectorizer:
    """TF-IDF 向量化器，基于词频统计"""

    def __init__(self):
        self.vocabulary: Dict[str, int] = {}
        self.idf: Dict[str, float] = {}
        self._fitted = False

    @staticmethod
    def tokenize(text: str) -> List[str]:
        """简单分词：中英文混合"""
        # 提取中文词语（2-4字）和英文单词
        tokens = []
        # 中文（2-4字词）
        cn_words = re.findall(r'[\u4e00-\u9fff]{2,4}', text)
        tokens.extend(cn_words)
        # 英文单词
        en_words = re.findall(r'[a-zA-Z]+', text.lower())
        tokens.extend(en_words)
        # 数字
        nums = re.findall(r'\d+[\.\d]*', text)
        tokens.extend(nums)
        return tokens

    def fit(self, documents: List[str]) -> 'TfidfVectorizer':
        """构建词汇表和 IDF 值"""
        n_docs = len(documents)
        df: Counter = Counter()

        for doc in documents:
            tokens = set(self.tokenize(doc))
            for token in tokens:
                df[token] += 1

        vocab_list = sorted(df.keys())
        self.vocabulary = {word: idx for idx, word in enumerate(vocab_list)}
        
        for word, freq in df.items():
            self.idf[word] = math.log((n_docs + 1) / (freq + 1)) + 1
        
        self._fitted = True
        return self

    def transform(self, documents: List[str]) -> list:
        """将文档转换为 TF-IDF 向量"""
        import numpy as np
        n_docs = len(documents)
        dim = len(self.vocabulary)
        
        vectors = np.zeros((n_docs, dim), dtype=np.float32)
        
        for i, doc in enumerate(documents):
            tokens = self.tokenize(doc)
            tf: Counter = Counter(tokens)
            
            for token, count in tf.items():
                if token in self.vocabulary:
                    idx = self.vocabulary[token]
                    vectors[i, idx] = (1 + math.log(count)) * self.idf.get(token, 1.0)
            
            # L2 归一化
            norm = np.linalg.norm(vectors[i])
            if norm > 0:
                vectors[i] /= norm
        
        return vectors.tolist()

    def fit_transform(self, documents: List[str]) -> list:
        self.fit(documents)
        return self.transform(documents)


# ══════════════════════════════════════════════════
# 5. 纯 NumPy 向量数据库
# ══════════════════════════════════════════════════

class SimpleVectorDB:
    """轻量级向量数据库，使用 NumPy + JSON 持久化"""

    def __init__(self, db_path: str, collection_name: str):
        self.db_path = Path(db_path)
        self.collection_name = collection_name
        self.col_dir = self.db_path / collection_name
        self.vectors: Optional[list] = None  # list of float lists
        self.ids: List[str] = []
        self.documents: List[str] = []
        self.metadatas: List[Dict] = []

    @property
    def count(self) -> int:
        return len(self.ids)

    def _save(self):
        """保存到磁盘"""
        self.col_dir.mkdir(parents=True, exist_ok=True)
        
        import numpy as np
        
        # 保存向量
        if self.vectors:
            np.save(self.col_dir / "vectors.npy", np.array(self.vectors, dtype=np.float32))
        
        # 保存元数据
        data = {
            "ids": self.ids,
            "documents": self.documents,
            "metadatas": self.metadatas,
            "collection_name": self.collection_name,
        }
        with open(self.col_dir / "metadata.json", 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def _load(self) -> bool:
        """从磁盘加载"""
        meta_file = self.col_dir / "metadata.json"
        vec_file = self.col_dir / "vectors.npy"
        
        if not meta_file.exists() or not vec_file.exists():
            return False
        
        import numpy as np
        
        with open(meta_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        self.ids = data["ids"]
        self.documents = data["documents"]
        self.metadatas = data["metadatas"]
        self.vectors = np.load(vec_file).tolist()
        
        return True

    def add(self, ids: List[str], documents: List[str],
            metadatas: List[Dict], embeddings: list):
        """添加记录"""
        if self.vectors is None:
            self.vectors = []
        
        self.ids.extend(ids)
        self.documents.extend(documents)
        self.metadatas.extend(metadatas)
        self.vectors.extend(embeddings)

    def query(self, query_embedding: list, top_k: int = 5,
              include: Optional[List[str]] = None) -> Dict:
        """查询最相似的文档"""
        import numpy as np
        
        if not self.vectors or len(self.vectors) == 0:
            return {"ids": [[]], "documents": [[]], "metadatas": [[]], "distances": [[]]}
        
        query_vec = np.array(query_embedding, dtype=np.float32).reshape(1, -1)
        all_vecs = np.array(self.vectors, dtype=np.float32)
        
        # 余弦相似度
        query_norm = np.linalg.norm(query_vec)
        if query_norm > 0:
            query_vec = query_vec / query_norm
        
        norms = np.linalg.norm(all_vecs, axis=1, keepdims=True)
        norms = np.where(norms > 0, norms, 1)
        normalized = all_vecs / norms
        
        similarities = np.dot(normalized, query_vec.T).flatten()
        
        # Top-k 索引（降序）
        k = min(top_k, len(similarities))
        top_indices = np.argsort(similarities)[::-1][:k]
        
        result_ids = [[str(self.ids[i]) for i in top_indices]]
        result_docs = [[self.documents[i] for i in top_indices]]
        result_metas = [[self.metadatas[i] for i in top_indices]]
        result_dists = [[float(1 - similarities[i]) for i in top_indices]]
        
        result = {
            "ids": result_ids,
            "documents": result_docs,
            "metadatas": result_metas,
            "distances": result_dists,
        }
        return result

    def get(self, limit: Optional[int] = None,
            include: Optional[List[str]] = None) -> Dict:
        """获取记录"""
        if limit is None:
            limit = len(self.ids)
        
        result = {
            "ids": self.ids[:limit],
            "documents": self.documents[:limit],
            "metadatas": self.metadatas[:limit],
        }
        return result

    def delete_collection(self):
        """删除集合数据"""
        if self.col_dir.exists():
            import shutil
            shutil.rmtree(self.col_dir)
        self.vectors = [] if self.vectors else None
        self.ids.clear()
        self.documents.clear()
        self.metadatas.clear()


# ══════════════════════════════════════════════════
# 主流程
# ══════════════════════════════════════════════════

def build_knowledge_base() -> Dict[str, Any]:
    """
    构建知识库主函数
    
    Returns:
        包含构建结果的字典
    """
    from langchain_core.documents import Document as LCDocument

    logger.info("=" * 50)
    logger.info("开始构建知识库")
    logger.info(f"   数据目录: {DATA_DIR}")
    logger.info(f"   输出目录: {VECTOR_DB_DIR}")
    logger.info("=" * 50)

    # ── Step 1: 扫描并加载文档 ──
    logger.info("\n[Step 1] 加载文档...")
    
    if not DATA_DIR.exists():
        logger.error(f"数据目录不存在: {DATA_DIR}")
        return {"status": "error", "message": f"数据目录不存在: {DATA_DIR}"}

    file_patterns = ["*.pdf", "*.docx", "*.doc"]
    files = []
    for pattern in file_patterns:
        files.extend(sorted(DATA_DIR.glob(pattern)))
    
    if not files:
        logger.warning("未找到任何文档文件!")
        return {"status": "error", "message": "未找到任何文档文件"}

    raw_docs: List[LCDocument] = []
    for filepath in files:
        fname = os.path.basename(filepath)
        logger.info(f"   处理: {fname}")

        raw_text = load_document(str(filepath))
        if not raw_text or len(raw_text.strip()) < 20:
            logger.warning(f"     跳过（内容为空或过短）")
            continue

        cleaned = clean_text(raw_text)
        raw_docs.append(LCDocument(page_content=cleaned, metadata={"source": fname}))
        logger.info(f"     OK ({len(cleaned)} 字符)")

    logger.info(f"   共加载 {len(raw_docs)} 个文档")

    # ── Step 2: 分块 ──
    logger.info("\n[Step 2] 文本分块...")
    chunks = split_documents(raw_docs)
    logger.info(f"   生成了 {len(chunks)} 个文本块")

    # ── Step 3: TF-IDF 向量化 ──
    logger.info("\n[Step 3] TF-IDF 向量化...")
    vectorizer = TfidfVectorizer()
    texts = [chunk.page_content for chunk in chunks]
    embeddings = vectorizer.fit_transform(texts)
    logger.info(f"   向量维度: {len(vectorizer.vocabulary)}")
    logger.info(f"   词汇量: {len(vectorizer.vocabulary)}")

    # ── Step 4: 存储到向量库 ──
    logger.info("\n[Step 4] 存储向量库...")
    db = SimpleVectorDB(str(VECTOR_DB_DIR), COLLECTION_NAME)

    # 删除旧数据
    db.delete_collection()

    ids = [f"chunk_{i:04d}" for i in range(len(chunks))]
    metadatas = [dict(chunk.metadata) for chunk in chunks]
    docs = [chunk.page_content for chunk in chunks]

    db.add(ids=ids, documents=docs, metadatas=metadatas, embeddings=embeddings)
    db._save()

    # 保存 vectorizer 以便后续查询时使用（使用 JSON 避免 pickle 兼容性问题）
    vec_data = {
        "vocabulary": vectorizer.vocabulary,
        "idf": {k: float(v) for k, v in vectorizer.idf.items()},
    }
    with open(db.col_dir / "vectorizer.json", 'w', encoding='utf-8') as f:
        json.dump(vec_data, f)

    logger.info(f"   已存储 {db.count} 条记录")

    # ── 完成 ──
    result = {
        "status": "success",
        "documents_loaded": len(raw_docs),
        "chunks_created": db.count,
        "vector_dim": len(vectorizer.vocabulary),
        "vocab_size": len(vectorizer.vocabulary),
        "collection_name": COLLECTION_NAME,
        "storage_path": str(db.col_dir),
    }

    logger.info("\n" + "=" * 50)
    logger.info("知识库构建完成!")
    for key, value in result.items():
        if key != "status":
            logger.info(f"   {key}: {value}")
    logger.info("=" * 50)

    return result


if __name__ == "__main__":
    import json
    result = build_knowledge_base()
    print(json.dumps(result, ensure_ascii=False, indent=2))
